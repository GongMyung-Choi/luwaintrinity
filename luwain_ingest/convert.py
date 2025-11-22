import os, sys, json
from pathlib import Path

def write_output(base, text):
    txt = Path(f"{base}.txt")
    txt.write_text(text, encoding="utf-8")
    return str(txt)

def read_docx(path):
    from docx import Document
    return "\n".join(p.text for p in Document(path).paragraphs)

def read_pdf(path):
    from pdfminer.high_level import extract_text
    return extract_text(path)

def read_image_ocr(path):
    import pytesseract
    from PIL import Image
    return pytesseract.image_to_string(Image.open(path), lang="kor+eng")

def read_hwp(path):
    import hwp5
    from hwp5.xmlmodel import HWP5File
    buf = []
    with open(path, "rb") as f:
        h = HWP5File(hwp5.HWP5File(f))
        for sec in h.bodytext():
            for p in sec.paragraphs():
                txt = "".join([c.text for c in p.texts()])
                if txt:
                    buf.append(txt)
    return "\n".join(buf)

def read_hwpx(path):
    from lxml import etree
    ns = {"hp": "http://www.hancom.co.kr/hwpml/2011/paragraph"}
    tree = etree.parse(path)
    return "\n".join(t.text for t in tree.xpath("//hp:t", namespaces=ns) if t.text)

def convert_any(path):
    p = Path(path)
    ext = p.suffix.lower()
    base = str(p.with_suffix(""))

    readers = {
        ".docx": read_docx,
        ".pdf": read_pdf,
        ".hwp": read_hwp,
        ".hwpx": read_hwpx,
    }

    try:
        if ext in readers:
            text = readers[ext](p)
        elif ext in [".png", ".jpg", ".jpeg", ".bmp"]:
            text = read_image_ocr(p)
        elif ext in [".txt", ".md"]:
            text = Path(p).read_text(encoding="utf-8", errors="ignore")
        else:
            return {"ok": False, "hint": f"미지원 확장자: {ext}"}
        out = write_output(base, text)
        return {"ok": True, "out": out}
    except Exception as e:
        return {"ok": False, "hint": str(e)}

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("사용법: python convert.py <파일경로>")
        sys.exit(1)
    print(json.dumps(convert_any(sys.argv[1]), ensure_ascii=False, indent=2))
